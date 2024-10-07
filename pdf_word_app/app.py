from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import os
import pdfplumber
from docx import Document
from docx.shared import Inches
import io
import sqlite3
from datetime import datetime

app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for flashing messages
UPLOAD_FOLDER = 'uploads'
VERSION_FOLDER = 'versions'
DB_PATH = 'versions.db'

# Create necessary folders if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(VERSION_FOLDER, exist_ok=True)

# SQLite setup for version control
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_name TEXT,
            timestamp TEXT
        )
    ''')
    conn.commit()
    conn.close()

# Initialize DB
init_db()

# Home Route
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if 'pdfFile' in request.files:
            # Upload and scrape PDF
            pdf_file = request.files['pdfFile']
            if pdf_file.filename != '':
                file_path = os.path.join(UPLOAD_FOLDER, pdf_file.filename)
                pdf_file.save(file_path)

                # Save scraped content (as text) into versions
                scrape_pdf_and_save(file_path)

                flash(f"PDF {pdf_file.filename} uploaded and scraped.", "success")
                return redirect(url_for('index'))
        elif 'docxFile' in request.files:
            # Upload edited Word document
            docx_file = request.files['docxFile']
            if docx_file.filename != '':
                docx_path = os.path.join(VERSION_FOLDER, docx_file.filename)
                docx_file.save(docx_path)

                # Save version to DB
                save_version_to_db(docx_file.filename)
                flash(f"Edited Word document {docx_file.filename} uploaded and saved.", "success")
                return redirect(url_for('index'))

    # Retrieve versions for display
    versions = get_versions_from_db()

    return render_template("upload.html", versions=versions)


# PDF scraping and save to versions folder (text, tables, images, formatting)
def scrape_pdf_and_save(file_path):
    doc = Document()  # Create a new Word document

    with pdfplumber.open(file_path) as pdf:
        # Iterate over each page
        for page_num, page in enumerate(pdf.pages):
            # Extract text and add to Word document
            text = page.extract_text()
            if text:
                doc.add_paragraph(f"Page {page_num + 1} Text:\n{text}")

            # Extract and format tables
            tables = page.extract_table()
            if tables:
                doc.add_paragraph(f"Page {page_num + 1} Tables:")
                for table in tables:
                    table_str = '\n'.join(['\t'.join(row) for row in table if row])
                    doc.add_paragraph(table_str)

            # Extract images and add to Word document
            for image in page.images:
                # Get image properties and extract it
                img_bbox = (image['x0'], image['top'], image['x1'], image['bottom'])
                img = page.within_bbox(img_bbox).to_image()

                # Save image to a bytes stream
                img_stream = io.BytesIO()
                img.save(img_stream, format="PNG")
                img_stream.seek(0)

                # Add image to the Word document
                doc.add_paragraph(f"Page {page_num + 1} Image:")
                doc.add_picture(img_stream, width=Inches(4.5))  # Adjust width as necessary

    # Save document to file
    output_filename = os.path.splitext(os.path.basename(file_path))[0] + "_scraped.docx"
    output_path = os.path.join(VERSION_FOLDER, output_filename)
    doc.save(output_path)

    # Save version to DB
    save_version_to_db(output_filename)


# Save version info in the database
def save_version_to_db(filename):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cursor.execute("INSERT INTO versions (file_name, timestamp) VALUES (?, ?)", (filename, timestamp))
    conn.commit()
    conn.close()


# Get versions from the database
def get_versions_from_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT id, file_name, timestamp FROM versions ORDER BY id DESC")
    versions = cursor.fetchall()
    conn.close()
    return [{"id": row[0], "file_name": row[1], "timestamp": row[2]} for row in versions]


# Generate Word document (from initial PDF)
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    # Assuming latest scraped PDF file is used for Word generation
    latest_version = get_versions_from_db()[0]
    return send_file(os.path.join(VERSION_FOLDER, latest_version['file_name']), as_attachment=True)


# Restore specific version
@app.route("/restore_version", methods=["POST"])
def restore_version():
    version_id = request.form['version_id']
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT file_name FROM versions WHERE id=?", (version_id,))
    version = cursor.fetchone()
    conn.close()

    if version:
        file_name = version[0]
        file_path = os.path.join(VERSION_FOLDER, file_name)
        return send_file(file_path, as_attachment=True)


# Delete specific version
@app.route("/delete_version", methods=["POST"])
def delete_version():
    version_id = request.form['version_id']
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM versions WHERE id=?", (version_id,))
    conn.commit()
    conn.close()

    flash("Version deleted successfully.", "success")
    return redirect(url_for('index'))


if __name__ == "__main__":
    app.run(debug=True)
